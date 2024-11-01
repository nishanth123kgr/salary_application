import smtplib
import sys
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import pandas as pd
from num2words import num2words
from docx import Document
# from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter
import os
import logging 
import concurrent.futures
import subprocess




#now we will Create and configure logger 
logging.basicConfig(filename="std.log", 
					format='%(asctime)s %(message)s', 
					filemode='w') 

#Let us Create an object 
logger=logging.getLogger() 

#Now we are going to Set the threshold of logger to DEBUG 
logger.setLevel(logging.DEBUG) 


class SalaryReport:
    def __init__(self, period):
        self.period = period
        self.base = 'C:/Salary_Application/salary_reports/'
        if not os.path.exists(self.base + 'pdf/encrypted'):
            os.makedirs(self.base + 'pdf/encrypted')
            
    def convert(self, docx_path, pdf_loc):
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"{docx_path} does not exist")
        
        # print(docx_path, pdf_loc)
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir', pdf_loc], check=True)
        
        print(f"Converted {docx_path} to PDF @ {pdf_loc}")

    def encrypt_pdf(self, pdf, password="!@#$%^&*()"):
        output_encrypted_pdf = self.base + "pdf/encrypted/" + pdf.split("/")[-1][:-4] + "_encrypted.pdf"

        output_pdf_file = open(pdf, "rb")
        output_pdf_reader = PdfReader(output_pdf_file)
        output_pdf_writer = PdfWriter()

        for page_num in range(len(output_pdf_reader.pages)):
            output_pdf_writer.add_page(output_pdf_reader.pages[page_num])

        output_pdf_writer.encrypt(password)

        with open(output_encrypted_pdf, "wb") as output_encrypted_pdf_file:
            output_pdf_writer.write(output_encrypted_pdf_file)
        output_pdf_file.close()
        print("Encrypted PDF created successfully!")
        return output_encrypted_pdf

    def send_mail(self, toaddr, password, name, pdf_loc):
        # Email details
        sender_email = 'noreply@auttvl.ac.in'
        receiver_email = toaddr
        subject = f'AURCT Salary Slip for the month of {self.period}'
        body = 'Password is first two letters of your name in lower-case, date and month of joining. If your name is "Dr. S. John Doe" and date of joining is "01.01.2021", then your password is "jo0101".'

        # Attachment details
        attachment_path = pdf_loc  # Replace with your file path

        # SMTP server configuration
        smtp_server = 'smtp.gmail.com'
        smtp_port = 587
        email_password = 'ydcznicxgwjraeni' 

        # Create the email message
        message = MIMEMultipart()
        message['From'] = sender_email
        message['To'] = receiver_email
        message['Subject'] = subject

        # Attach body to the email
        message.attach(MIMEText(body, 'plain'))

        # Attach the file
        attachment = open(attachment_path, 'rb')
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header(
            'Content-Disposition',
            f'attachment; filename= {attachment_path.split("/")[-1]}'
        )
        message.attach(part)

        # Establish a secure connection with the SMTP server
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            # Login to your email account
            server.login(sender_email, email_password)
            # Send email
            server.sendmail(sender_email, receiver_email, message.as_string())
            print("Email sent successfully! to", receiver_email)

    def convert_name(self, x):
        x = x.strip().replace(' ', '').lower().split('.')
        try:
            return x[3] if len(x) > 3 and len(x[2]) == 1 else x[2] if len(x) > 2 else x[1]
        except:
            print(x)
            exit(1)

    def generate_password(self, name, doj):
        password = f"{self.convert_name(name)[:2]}{doj.strip().replace('.', '')[:4]}"
        print(password, name)
        return password
    
    
    
        
    


    def gen_sal_report(self, row, index, socketio=None):
        # base_path = sys._MEIPASS
        base_path = os.getcwd()
        template_path = os.path.join(base_path, 'sal_template.docx')
        # Open the document using the adjusted path
        doc = Document(template_path)
        placeholders = {
            'name': row[0].strip(),
            'des': row[20],
            'bp': row[2],
            'ddd': row[3],
            'hra': row[4],
            'cca': row[5],
            'ma': row[6],
            'te': row[7],
            'cps': row[8],
            'ca': row[9],
            'it': row[10],
            'nhis': row[11],
            'fsf': row[12],
            'fa': row[13],
            'pt': row[-1],
            'mcp': row[15],
            'td': row[16],
            'np': f"Rs.{row[17]}/- (Rupees {num2words(row[17]).replace(',', '').replace(' and', '').replace('-', ' ').title()} only)",
            'sno': row[18],
            'dep': row[19],
            'doj': row[21],
            'sop': row[22],
            'bn': row[23],
            'ac': row[24],
            'gpf': '-',
            "gpfa": '-',
            'payslip': f'Pay Slip for the month of {self.period}'
        }
        for table in doc.tables:
            fixed_elems = []
            for rows in table.rows:
                for cell in rows.cells:
                    if cell.text not in fixed_elems:
                        fixed_elems.append(cell.text)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                text = run.text
                                if text in placeholders:
                                    # print(text, placeholders[text])
                                    run.text = placeholders[text]
        self.period = self.period.replace(' ', '-')
        file_name = f"{row[0].replace(' ', '')}{self.period}_{row[18]}" 
        doc.save(f"{self.base}{file_name}.docx")
        pdf_loc = f"{self.base}pdf/{file_name}.pdf"
        self.convert(f"{self.base}{file_name}.docx", f'{self.base}pdf/')
        output = self.encrypt_pdf(pdf_loc, self.generate_password(row[0], row[21]))
        socketio.emit('report-ready', index)
        # self.send_mail(row[-2], '', row[0], output)
        
        self.send_mail('nishanth123kgr@gmail.com', '', row[0], output)
        socketio.emit('mail-sent', index)



def read_data(salary, staff, month="Month 2XXX"):
    data = pd.read_excel(salary)
    # _, file = os.path.split(salary)
    period = month
    d1 = pd.read_excel(staff)
    df = data.iloc[2:, 1:]
    df.columns = df.iloc[0]
    df = df[1:]
    df = df.dropna()
    df.reset_index(drop=True, inplace=True)
    d1.columns = d1.iloc[1]
    d1 = d1.iloc[2:, 1:]
    d1 = d1.dropna()
    d1.reset_index(drop=True, inplace=True)
    df.insert(1, 'Staff_name', df['Name'])
    report = SalaryReport(period)
    # d1["Name"] = d1["Name"].apply(lambda x: report.convert_name(x))
    # df["Name"] = df["Name"].apply(lambda x: report.convert_name(x))
    # print(d1)
    merged = pd.merge(df, d1, on='Employee_No', how='inner')
    print(merged.columns)
    merged = merged.drop(['Name_x', 'Name_y'], axis=1)
    
    # exit()
    merged.reset_index(drop=True, inplace=True)
    # print(merged)

    merged['E-mail'] = merged['E-mail'].apply(lambda x: x.replace(' ', ''))

    # merged = merged.iloc[11:13, :]

    print(merged)

    return merged, report

def generate_reports(data, report, socketio=None):
    for index, row in data.iterrows():
        row = row.astype(str)
        row = row.to_list()
        # print(row)
        row.append(row.pop(16))
        print(row)
        
        try:
            report.gen_sal_report(row, index, socketio)
            logger.info('Email successfully sent to '+ row[0]+ ' Index: '+str(index)) 
            
            # exit()
        except Exception as e:
            logger.error(f"An exception occurred for index {index}: {e}")
            print("Error at index", index)



# def generate_report_for_row(row, index, report, socketio):
#     try:
#         row = row.astype(str).to_list()
#         row.append(row.pop(16))
#         report.gen_sal_report(row, index, socketio)
#         print("Created PDF", row[0])
#         logger.info(f'Email successfully sent to {row[0]} Index: {index}')
#     except Exception as e:
#         logger.error(f"Error at index {index}: {e}")
#         return False
#     return True

# def generate_reports(data, report, socketio=None, max_workers=4):
#     with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
#         # Prepare arguments for each row
#         futures = [
#             executor.submit(generate_report_for_row, row, index, report, socketio)
#             for index, row in data.iterrows()
#         ]
        
#         for future in concurrent.futures.as_completed(futures):
#             if not future.result():
#                 logger.warning("A report generation task encountered an error.")

# def generate_report_for_row(row, index, report, socketio):
#     try:
#         # Convert the row to a list of strings and rearrange as needed
#         row = row.astype(str).to_list()
#         row.append(row.pop(16))  # Move the 17th element to the end
        
#         # Generate the salary report
#         report.gen_sal_report(row, index, socketio)
#         logger.info(f'Email successfully sent to {row[0]} Index: {index}')
        
#     except Exception as e:
#         logger.error(f"Error at index {index}: {e}")
#         return False
#     return True

# def generate_reports(data, report, socketio=None, max_workers=4):
#     # Ensure that data is a DataFrame
#     if not isinstance(data, pd.DataFrame):
#         logger.error("Data is not a pandas DataFrame.")
#         return
    
#     with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
#         # Prepare futures for each row
#         futures = {
#             executor.submit(generate_report_for_row, row, index, report, socketio): index
#             for index, row in data.iterrows()
#         }

#         for future in concurrent.futures.as_completed(futures):
#             index = futures[future]
#             try:
#                 if not future.result():
#                     logger.warning(f"Report generation failed for index {index}.")
#             except Exception as e:
#                 

if __name__ == '__main__':
    salary = 'Salary for the month August  2024 (Scale) - General Fund (Scale).xls'
    staff = 'Teaching Staff Salary -20231.xlsx'
    data, report = read_data(salary, staff, "August 2024")
    report.convert_to_pdf()
    exit(1)
    print(data)
    # data=pd.DataFrame(data.iloc[21, :]).T
    # exit()
    for index, row in data.iterrows():
        row = row.astype(str)
        row = row.to_list()
        # print(row)
        row.append(row.pop(16))
        print(row)
        
        try:
            report.gen_sal_report(row)
            logger.info('Email successfully sent to '+ row[0]+ ' Index: '+str(index)) 
            exit()
        except Exception as e:
            print(e)
            print("Error at index", index)
            logger.error(f"Error at index {index} for {row[0]}: {e}", exc_info=True)

