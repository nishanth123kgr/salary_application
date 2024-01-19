import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import pandas as pd
from num2words import num2words
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter
import os

class SalaryReport:
    def __init__(self, period):
        self.period = period

    def encrypt_pdf(self, pdf, password="!@#$%^&*()"):
        output_encrypted_pdf = "salary_reports/pdf/encrypted/" + pdf.split("/")[-1][:-4] + "_encrypted.pdf"

        output_pdf_file = open(pdf, "rb")
        output_pdf_reader = PdfReader(output_pdf_file)
        output_pdf_writer = PdfWriter()

        for page_num in range(len(output_pdf_reader.pages)):
            output_pdf_writer.add_page(output_pdf_reader.pages[page_num])

        output_pdf_writer.encrypt(password)

        with open(output_encrypted_pdf, "wb") as output_encrypted_pdf_file:
            output_pdf_writer.write(output_encrypted_pdf_file)
        output_pdf_file.close()
        print("Encrypted PDF created successfully!")
        return output_encrypted_pdf

    def send_mail(self, toaddr, password, name, pdf_loc):
        # Email details
        sender_email = 'noreply@auttvl.ac.in'
        receiver_email = toaddr
        subject = f'AURCT Salary Slip for the month of {self.period}'
        body = 'Password is first two letters of your name in lower-case, date and month of joining. If your name is "Dr. S. John Doe" and date of joining is "01.01.2021", then your password is "jo0101".'

        # Attachment details
        attachment_path = pdf_loc  # Replace with your file path

        # SMTP server configuration
        smtp_server = 'smtp.gmail.com'
        smtp_port = 587
        email_password = 'iprzeridukqaklyz'

        # Create the email message
        message = MIMEMultipart()
        message['From'] = sender_email
        message['To'] = receiver_email
        message['Subject'] = subject

        # Attach body to the email
        message.attach(MIMEText(body, 'plain'))

        # Attach the file
        attachment = open(attachment_path, 'rb')
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header(
            'Content-Disposition',
            f'attachment; filename= {attachment_path.split("/")[-1]}'
        )
        message.attach(part)

        # Establish a secure connection with the SMTP server
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            # Login to your email account
            server.login(sender_email, email_password)
            # Send email
            server.sendmail(sender_email, receiver_email, message.as_string())
            print("Email sent successfully!")
    def convert_name(self, x):
        x = x.strip().replace(' ', '').lower().split('.')
        try:
            return x[3] if len(x) > 3 and len(x[2])==1 else x[2] if len(x) > 2 else x[1] 
        except:
            print(x)
            exit(1)

    def generate_password(self, name, doj):
        password = f"{self.convert_name(name)[:2]}{doj.replace('.', '')[:4]}"
        print(password, name)
        return password


    def gen_sal_report(self, row):
        doc = Document("sal_template.docx")
        placeholders = {
            'name': row[0].strip(),
            'des': row[20],
            'bp': row[2],
            'ddd': row[3],
            'hra': row[4],
            'cca': row[5],
            'ma': row[6],
            'te': row[7],
            'cps': row[8],
            'ca': row[9],
            'it': row[10],
            'nhis': row[11],
            'fsf': row[12],
            'fa': row[13],
            'pt': row[14],
            'mcp': row[15],
            'td': row[16],
            'np': f"Rs.{row[17]}/- (Rupees {num2words(row[17]).replace(',', '').replace(' and', '').replace('-', ' ').title()} only)",
            'sno': row[18],
            'dep': row[19],
            'doj': row[21],
            'sop': row[22],
            'bn': row[23],
            'ac': row[24],
            'gpf': '-',
            "gpfa": '-',
            'payslip':f'Pay Slip for the month of {self.period}'
        }
        for table in doc.tables:
            fixed_elems = []
            for rows in table.rows:
                for cell in rows.cells:
                    if cell.text not in fixed_elems:
                        fixed_elems.append(cell.text)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                text = run.text
                                if text in placeholders:
                                    # print(text, placeholders[text])
                                    run.text = placeholders[text]
        doc.save(f"salary_reports/{row[0].replace(' ', '')}{self.period}.docx")
        pdf_loc = f"salary_reports/pdf/{row[0].replace(' ', '')}{self.period}.pdf"
        convert(f"salary_reports/{row[0].replace(' ', '')}{self.period}.docx", pdf_loc)
        output = self.encrypt_pdf(pdf_loc, self.generate_password(row[0], row[21]))
        # self.send_mail(row[-1], '', row[0], output)


def read_data(salary, staff):
    data = pd.read_excel(salary)
    print(data)
    file = salary.filename
    period = ' '.join(file.split()[5:7])
    d1 = pd.read_excel(staff)
    df = data.iloc[2:, 1:]
    print(df)
    df.columns = df.iloc[0]
    df = df[1:]
    df = df.dropna()
    df.reset_index(drop=True, inplace=True)
    d1.columns = d1.iloc[1]
    d1 = d1.iloc[2:, 1:]
    d1 = d1.dropna()
    d1.reset_index(drop=True, inplace=True)
    df.insert(1, 'Staff_name', df['Name'])
    report = SalaryReport(period)
    d1["Name"] = d1["Name"].apply(lambda x: report.convert_name(x))
    df["Name"] = df["Name"].apply(lambda x: report.convert_name(x))
    # print(d1)
    merged = pd.merge(df, d1, on='Name', how='inner')
    merged = merged.drop(['Name'], axis=1)
    merged.reset_index(drop=True, inplace=True)
    # print(merged)
    
    merged['E-mail'] = merged['E-mail'].apply(lambda x: x.replace(' ', ''))
    
    
    return merged, report

def generate_reports(data, report, socketio=None):
    for index, row in data.iterrows():
        row = row.astype(str)
        row = row.to_list()
        print(row)
        try:
            report.gen_sal_report(row)
            socketio.emit('report-ready', index)
        except Exception as e:
            print(e)
            print("Error at index", index)
            exit(1)
    
if __name__ == '__main__':
    for file in os.listdir('salary_datas'):
        merged, report = read_data('salary_datas/'+file, 'Teaching Staff Salary -20231.xlsx')
        generate_reports(merged, report)
        
        
        

